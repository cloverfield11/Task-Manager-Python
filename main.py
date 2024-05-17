from flask import Flask, render_template, request, redirect, url_for, jsonify, session, send_file
from werkzeug.utils import secure_filename
import os
import pandas as pd
import json
from datetime import datetime
from flask_sqlalchemy import SQLAlchemy
from io import BytesIO
from docx import Document
import openpyxl
from flask import jsonify

app = Flask(__name__)
app.secret_key = '54FGHhfd356h56hgh7H'
tasks_df = pd.read_excel('tasks.xlsx')
tasks_df['Progress'] = tasks_df['Progress'].apply(lambda x: [] if pd.isna(x) else x)
editing_task = None

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(80), nullable=False)

with app.app_context():
    db.create_all()
  
@app.route('/')
def welcome():
    if 'username' in session:
        return redirect(url_for('index'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        user = User.query.filter_by(username=username, password=password).first()

        if user:
            session['username'] = username
            return redirect(url_for('index'))

        return render_template('login.html', error='Invalid username or password')

    return render_template('login.html', error=None)

@app.route('/archive_tasks', methods=['POST'])
def archive_tasks():
    global tasks_df

    completed_task_indices = tasks_df.index[tasks_df['StageTask'] == 'Completed'].tolist()

    tasks_df.loc[completed_task_indices, 'StageTask'] = 'Archived'
    tasks_df.to_excel('tasks.xlsx', index=False)

    return jsonify({'status': 'success', 'message': 'Tasks archived successfully'})
  
@app.route('/move_to_archive', methods=['POST'])
def move_to_archive():
    global tasks_df

    archived_task_indices = tasks_df.index[tasks_df['StageTask'] == 'Archived'].tolist()

    archived_tasks_df = tasks_df.loc[archived_task_indices].copy()

    tasks_df.drop(index=archived_task_indices, inplace=True)

    tasks_df.to_excel('tasks.xlsx', index=False)

    if os.path.exists('archived_tasks.xlsx'):
        existing_archived_tasks_df = pd.read_excel('archived_tasks.xlsx')
        existing_archived_tasks_df = pd.concat([existing_archived_tasks_df, archived_tasks_df], ignore_index=True)
        existing_archived_tasks_df.to_excel('archived_tasks.xlsx', index=False)
    else:
        archived_tasks_df.to_excel('archived_tasks.xlsx', index=False)

    return jsonify({'status': 'success', 'message': 'Tasks moved to archive successfully'})

@app.route('/archive_index')
def archive_index():
    tasks_df = pd.read_excel('archived_tasks.xlsx')
    main_tasks = []
    additional_tasks = []
    archived_tasks = []

    russian_months = {
        1: 'january',
        2: 'february',
        3: 'march',
        4: 'april',
        5: 'may',
        6: 'june',
        7: 'july',
        8: 'august',
        9: 'september',
        10: 'october',
        11: 'november',
        12: 'december'
    }

    for _, task in tasks_df.iterrows():
        category = task['CatTask']
        status = task['StageTask']

        color = ''
        if status == 'Archived':
            color = 'gray'
        else:
            priority = task['Priority']
            if priority == 'High':
                color = 'red'
            elif priority == 'Medium':
                color = 'orange'
            elif priority == 'Low':
                color = 'green'

        formatted_due_date = datetime.strptime(task['Deadline'], '%Y-%m-%d').strftime('%d %m %Y')

        day, month, year = formatted_due_date.split()
        month = month.lstrip('0')
        month_key = int(month)
        formatted_due_date = f"{day} {russian_months.get(month_key, month)} {year}"

        task_info = {
            'task_id': len(main_tasks) + len(additional_tasks) + len(archived_tasks),
            'task': task.to_dict(),
            'color': color,
            'formatted_due_date': formatted_due_date
        }

        if status == 'Completed':
            archived_tasks.append(task_info)
        elif category == 'Current tasks':
            main_tasks.append(task_info)
        elif category == 'Additional tasks':
            additional_tasks.append(task_info)

    if 'username' in session:
        return render_template('archive_index.html',
                               main_tasks=main_tasks,
                               additional_tasks=additional_tasks,
                               completed_tasks=archived_tasks)
    else:
        return redirect(url_for('welcome'))
      
@app.route('/export_docx')
def export_docx():
    tasks = load_tasks_from_excel()

    doc = Document()

    doc.add_heading('Completed tasks', level=1)

    for task in tasks:
        if task.get('StageTask', '') == 'Completed':
            doc.add_heading(task['TaskName'], level=2)
            doc.add_paragraph(f"Description: {task['DescrTask']}")
            doc.add_paragraph(f"Category: {task['CatTask']}")
            doc.add_paragraph(f"Deadline: {task['Deadline']}")
            doc.add_paragraph(f"Comment: {task['MyComment']}")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name='completed_tasks.docx')

def load_tasks_from_excel():
  wb = openpyxl.load_workbook('tasks.xlsx')
  sheet = wb['Sheet1']

  tasks = []

  headers = [cell.value for cell in sheet[1]]

  for row in sheet.iter_rows(min_row=2, values_only=True):
      task = dict(zip(headers, row))
      tasks.append(task)

  return tasks

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            return render_template('register.html', error='Username already exists')

        new_user = User(username=username, password=password)
        db.session.add(new_user)
        db.session.commit()

        session['username'] = username
        return redirect(url_for('index'))

    return render_template('register.html', error=None)

@app.route('/logout', methods=['POST'])
def logout():
    session.pop('username', None)
    return redirect(url_for('welcome'))

@app.route('/index')
def index():
    main_tasks = []
    additional_tasks = []
    completed_tasks = []

    russian_months = {
        1: 'january',
        2: 'february',
        3: 'march',
        4: 'april',
        5: 'may',
        6: 'june',
        7: 'july',
        8: 'august',
        9: 'september',
        10: 'october',
        11: 'november',
        12: 'december'
    }
  
    for _, task in tasks_df.iterrows():
        category = task['CatTask']
        status = task['StageTask']

        color = ''
        if status == 'Completed':
            color = 'gray'
        else:
            priority = task['Priority']
            if priority == 'High':
                color = 'red'
            elif priority == 'Medium':
                color = 'orange'
            elif priority == 'Low':
                color = 'green'

        formatted_due_date = datetime.strptime(task['Deadline'], '%Y-%m-%d').strftime('%d %m %Y')

        day, month, year = formatted_due_date.split()
        month = month.lstrip('0')
        month_key = int(month)
        formatted_due_date = f"{day} {russian_months.get(month_key, month)} {year}"
          
        task_info = {
            'task_id': len(main_tasks) + len(additional_tasks) + len(completed_tasks),
            'task': task.to_dict(),
            'color': color,
            'formatted_due_date': formatted_due_date
        }

        if status == 'Completed':
            completed_tasks.append(task_info)
        elif category == 'Current tasks':
            main_tasks.append(task_info)
        elif category == 'Additional tasks':
            additional_tasks.append(task_info)

    if 'username' in session:
        return render_template('index.html',
                               main_tasks=main_tasks,
                               additional_tasks=additional_tasks,
                               completed_tasks=completed_tasks)
    else:
        return redirect(url_for('welcome'))
      
@app.template_filter('to_date')
def to_date(value):
  return datetime.strptime(value, "%Y-%m-%d").strftime("%d.%m.%Y")

  
@app.route('/add_task', methods=['GET', 'POST'])
def add_task():
    global tasks_df

    if request.method == 'POST':
        new_task = {
            'TaskName': request.form['title'],
            'DescrTask': request.form['description'],
            'CatTask': request.form['category'],
            'Deadline': request.form['deadline'],
            'StageTask': request.form['status'],
            'MyComment': request.form['my_comment'],
            'Priority': request.form['priority'],
            'YourComment': '',
            'filename': request.form['filename'],
            'Progress': []
        }
        tasks_df = pd.concat([tasks_df, pd.DataFrame([new_task])],
                             ignore_index=True)
        tasks_df.to_excel('tasks.xlsx', index=False)

        return redirect(url_for('index'))

    if 'username' in session:
      return render_template('add_task.html')
    else:
      return redirect(url_for('welcome'))

@app.route('/task/<int:task_id>', methods=['GET', 'POST'])
def task(task_id):
    global editing_task

    if request.method == 'POST':
        if 'your_comment' in request.form:
            tasks_df.at[task_id, 'YourComment'] = request.form['your_comment']
        elif 'my_comment' in request.form:
            tasks_df.at[task_id, 'MyComment'] = request.form['my_comment']
        elif 'new_event' in request.form:
            new_event = {
                'Дата': datetime.now().strftime("%d.%m.%Y"),
                'Событие': request.form['new_event']
            }
            events_list = json.loads(tasks_df.at[task_id, 'Progress']) if isinstance(tasks_df.at[task_id, 'Progress'], str) else []
            events_list.append(new_event)

            tasks_df.at[task_id, 'Progress'] = json.dumps(events_list)

            tasks_df.to_excel('tasks.xlsx', index=False)

        

    if editing_task is not None:
        task = editing_task
        editing_task = None  
    else:
        task = tasks_df.iloc[task_id].copy()

    task['Deadline'] = datetime.strptime(task['Deadline'], "%Y-%m-%d").strftime("%d.%m.%Y")

    if 'Progress' in task:
        events_str = task['Progress']
        if isinstance(events_str, str):
            try:
                task['Progress'] = json.loads(events_str)
            except json.JSONDecodeError as e:
                print(f"Error decoding 'Progress': {e}")
                task['Progress'] = []

    if 'username' in session:
      return render_template('task.html', task=task, task_id=task_id)
    else:
      return redirect(url_for('welcome'))
  

@app.route('/add_event/<int:task_id>', methods=['POST'])
def add_event(task_id):
    new_event_description = request.form.get('new_event')

    current_date = datetime.now().strftime("%d.%m.%Y")

    events_list = tasks_df.at[task_id, 'Progress']

    events_list = eval(events_list) if isinstance(events_list, str) else events_list

    events_list.append({'date': current_date, 'description': new_event_description})

    events_json = json.dumps(events_list)

    tasks_df.at[task_id, 'Progress'] = events_json

    tasks_df.to_excel('tasks.xlsx', index=False)

    return redirect(url_for('task', task_id=task_id))
  
@app.route('/edit_task/<int:task_id>', methods=['GET', 'POST'])
def edit_task(task_id):
  task = tasks_df.iloc[task_id]
  if request.method == 'POST':
    tasks_df.at[task_id, 'TaskName'] = request.form['title']
    tasks_df.at[task_id, 'DescrTask'] = request.form['description']
    tasks_df.at[task_id, 'CatTask'] = request.form['category']
    tasks_df.at[task_id, 'Deadline'] = request.form['deadline']
    tasks_df.at[task_id, 'StageTask'] = request.form['status']
    tasks_df.at[task_id, 'MyComment'] = request.form['my_comment']
    tasks_df.at[task_id, 'Priority'] = request.form['priority']
    tasks_df.at[task_id, 'filename'] = request.form['filename']
    tasks_df.to_excel('tasks.xlsx', index=False)
    return redirect(url_for('task', task_id=task_id))

  if 'username' in session:
    return render_template('edit_task.html', task=task, task_id=task_id)
  else:
    return redirect(url_for('welcome'))

@app.route('/delete_task/<int:task_id>', methods=['POST'])
def delete_task(task_id):
  global tasks_df

  tasks_df = tasks_df.drop(task_id).reset_index(drop=True)
  tasks_df.to_excel('tasks.xlsx', index=False)

  return redirect(url_for('index'))

@app.route('/error/no_link', methods=['GET'])
def no_link_error():
    return render_template('no_link_error.html')
  
if __name__ == '__main__':
  app.run(host='0.0.0.0', port=5000)
